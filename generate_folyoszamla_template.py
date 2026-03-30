"""
Folyószámla összesítő Word sablon generátor.

Futtatás:
    pip install python-docx
    python generate_folyoszamla_template.py

Kimenet: folyoszamla_osszesito_template.docx
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        tblBorders.append(border)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def set_cell_paragraph(cell, text, bold=False, font_size=None,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Clear a cell and add a single run with the given formatting."""
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = alignment
    run = para.add_run(text)
    run.bold = bold
    if font_size is not None:
        run.font.size = Pt(font_size)
    return run


def main():
    doc = Document()

    # ------------------------------------------------------------------ #
    # Page setup: landscape A4, 1 cm margins                              #
    # ------------------------------------------------------------------ #
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)

    # Default style: Arial 8 pt
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(8)

    # ------------------------------------------------------------------ #
    # Header table (2 rows × 3 cols, no borders)                          #
    # ------------------------------------------------------------------ #
    header_table = doc.add_table(rows=2, cols=3)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(header_table)

    # Row 0: logo | title | date
    set_cell_paragraph(header_table.cell(0, 0), "{{LOGO}}",
                       alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_paragraph(header_table.cell(0, 1), "Folyószámla összesítő",
                       bold=True, font_size=14,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_paragraph(header_table.cell(0, 2), "Dátum: {{DATUM}}",
                       alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # Row 1: empty | year/month | page
    header_table.cell(1, 0).paragraphs[0].clear()
    set_cell_paragraph(header_table.cell(1, 1), "{{EV}}. {{HONAP}}",
                       bold=True, font_size=12,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_paragraph(header_table.cell(1, 2), "Oldal: {{OLDAL}}",
                       alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.add_paragraph("")  # spacer

    # ------------------------------------------------------------------ #
    # Main data table (Table Grid style, 11 columns)                      #
    # ------------------------------------------------------------------ #
    NUM_COLS = 11
    # rows: 0=group headers, 1=column headers, 2=prev-month balance,
    #       3..22=data rows (20), 23=comment row
    main_table = doc.add_table(rows=3, cols=NUM_COLS)
    main_table.style = "Table Grid"
    main_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ---- Row 0: group header row ------------------------------------ #
    r0 = main_table.rows[0]

    # Col 0: empty (Banki dátum column has no group)
    r0.cells[0].paragraphs[0].clear()

    # Cols 1–3: "Jóváírások"
    r0.cells[1].merge(r0.cells[3])
    set_cell_paragraph(r0.cells[1], "Jóváírások",
                       bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Cols 4–5: "Terhelések"
    r0.cells[4].merge(r0.cells[5])
    set_cell_paragraph(r0.cells[4], "Terhelések",
                       bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Cols 6–7: "Jogi költségek"
    r0.cells[6].merge(r0.cells[7])
    set_cell_paragraph(r0.cells[6], "Jogi költségek",
                       bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Col 8: "Banki kezelési költség"
    set_cell_paragraph(r0.cells[8], "Banki kezelési költség",
                       bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Col 9: "Egyenleg leemelés"
    set_cell_paragraph(r0.cells[9], "Egyenleg leemelés",
                       bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Col 10: "Záróegyenleg"
    set_cell_paragraph(r0.cells[10], "Záróegyenleg",
                       bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # ---- Row 1: column sub-headers ---------------------------------- #
    column_headers = [
        "Banki dátum",
        "Számla bevét",
        "Áttutóra bevét",
        "Egyéb",
        "Számláról túlfiz\nvisszautalás",
        "Áttutóról\nvisszautalás",
        "HM VGH jogi\nköltség",
        "HM EI jogi\nköltség",
        "Banki kezelési\nköltség",
        "Egyenleg\nleemelés",
        "Záróegyenleg",
    ]
    r1 = main_table.rows[1]
    for i, header_text in enumerate(column_headers):
        set_cell_paragraph(r1.cells[i], header_text,
                           bold=True, font_size=7,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # ---- Row 2: previous month closing balance ---------------------- #
    r2 = main_table.rows[2]
    # Cols 0–9 remain empty
    for i in range(10):
        r2.cells[i].paragraphs[0].clear()
    prev_text = (
        "{{EV_ELOZO}}. {{HONAP_ELOZO}}. havi záróegyenleg:\n"
        "{{ELOZO_HAVI_ZAROEGYENLEG}}"
    )
    set_cell_paragraph(r2.cells[10], prev_text,
                       font_size=7, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # ---- Rows 3–22: 20 placeholder data rows ------------------------ #
    data_placeholders = [
        ("{{BANKI_DATUM_%d}}", WD_ALIGN_PARAGRAPH.LEFT),
        ("{{SZAMLA_BEVET_%d}}", WD_ALIGN_PARAGRAPH.RIGHT),
        ("{{ATTUTORA_BEVET_%d}}", WD_ALIGN_PARAGRAPH.RIGHT),
        ("{{EGYEB_JOVAIRAS_%d}}", WD_ALIGN_PARAGRAPH.RIGHT),
        ("{{SZAMLAROL_TULFIZ_%d}}", WD_ALIGN_PARAGRAPH.RIGHT),
        ("{{ATTUTOR_VISSZAUT_%d}}", WD_ALIGN_PARAGRAPH.RIGHT),
        ("{{HM_VGH_JOGI_%d}}", WD_ALIGN_PARAGRAPH.RIGHT),
        ("{{HM_EI_JOGI_%d}}", WD_ALIGN_PARAGRAPH.RIGHT),
        ("{{BANKI_KEZ_KOLTSEG_%d}}", WD_ALIGN_PARAGRAPH.RIGHT),
        ("{{EGYENLEG_LEEMELES_%d}}", WD_ALIGN_PARAGRAPH.RIGHT),
        ("{{ZAROEGYENLEG_%d}}", WD_ALIGN_PARAGRAPH.RIGHT),
    ]

    for n in range(1, 21):
        row = main_table.add_row()
        for col_idx, (placeholder_fmt, align) in enumerate(data_placeholders):
            set_cell_paragraph(row.cells[col_idx],
                               placeholder_fmt % n,
                               font_size=7, alignment=align)

    # ---- Last row: merged comment/note ------------------------------ #
    note_row = main_table.add_row()
    note_row.cells[0].merge(note_row.cells[NUM_COLS - 1])
    set_cell_paragraph(note_row.cells[0], "{{MEGJEGYZES}}",
                       alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # ------------------------------------------------------------------ #
    # Column widths                                                        #
    # ------------------------------------------------------------------ #
    col_widths = [
        Cm(2.5),  # Banki dátum
        Cm(2.2),  # Számla bevét
        Cm(2.2),  # Áttutóra bevét
        Cm(2.2),  # Egyéb
        Cm(2.2),  # Számláról túlfiz visszautalás
        Cm(2.2),  # Áttutóról visszautalás
        Cm(2.2),  # HM VGH jogi költség
        Cm(2.2),  # HM EI jogi költség
        Cm(2.2),  # Banki kezelési költség
        Cm(2.2),  # Egyenleg leemelés
        Cm(3.5),  # Záróegyenleg
    ]
    for row in main_table.rows:
        for idx, width in enumerate(col_widths):
            # After merges some cells may share a tc element; set width on each
            cell = row.cells[idx]
            cell.width = width

    # ------------------------------------------------------------------ #
    # Save                                                                 #
    # ------------------------------------------------------------------ #
    output_filename = "folyoszamla_osszesito_template.docx"
    doc.save(output_filename)
    print(f"Template saved as: {output_filename}")


if __name__ == "__main__":
    main()
